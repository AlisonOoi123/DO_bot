from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        val = kwargs.get(edge, 'single')
        sz = kwargs.get(f'{edge}_sz', 4)
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), val)
        tag.set(qn('w:sz'), str(sz))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), '000000')
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_bullet(cell, text, font_size=10):
    p = cell.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    return p

def make_activity_table(doc, weeks_data):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0].cells
    hdr[0].width = Inches(0.8)
    hdr[1].width = Inches(5.5)
    for cell in hdr:
        set_cell_bg(cell, 'D3D3D3')
    p0 = hdr[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run0 = p0.add_run('Week')
    run0.bold = True
    run0.font.size = Pt(11)
    run0.font.name = 'Times New Roman'
    p1 = hdr[1].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run('Projects / Activities')
    run1.bold = True
    run1.font.size = Pt(11)
    run1.font.name = 'Times New Roman'

    for week_num, bullets in weeks_data:
        row = table.add_row()
        row.cells[0].width = Inches(0.8)
        row.cells[1].width = Inches(5.5)
        # Week number cell
        wc = row.cells[0]
        wc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        wp = wc.paragraphs[0]
        wp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        wr = wp.add_run(str(week_num))
        wr.bold = True
        wr.font.size = Pt(11)
        wr.font.name = 'Times New Roman'
        # Activities cell - clear default paragraph first
        ac = row.cells[1]
        ac.paragraphs[0]._element.getparent().remove(ac.paragraphs[0]._element)
        for bullet_text in bullets:
            add_bullet(ac, bullet_text)

    return table

def build_report(doc, month_year, weeks_data, trainee_date, supervisor_date):
    # Header
    for line, bold, size in [
        ('Tunku Abdul Rahman University of', True, 13),
        ('Management and Technology', True, 13),
        ('Faculty of Computing and Information Technology', True, 13),
        ('Industrial Training Progress Report', True, 13),
        ('Activity Log', False, 12),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Info fields
    for label, value in [
        ('Name of Trainee:', 'OOI YU ZHEN'),
        ('Name of Company:', 'Eng Sheng Sdn. Bhd.'),
        ('Month/Year', month_year),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        r_label = p.add_run(label)
        r_label.font.size = Pt(11)
        r_label.font.name = 'Times New Roman'
        p.add_run('\t\t')
        r_val = p.add_run(value)
        r_val.font.size = Pt(11)
        r_val.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Activity table
    make_activity_table(doc, weeks_data)

    doc.add_paragraph()

    # Suggestions box
    p = doc.add_paragraph()
    r = p.add_run('Suggestions / Comments / Additional information (if any):')
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    sug_table = doc.add_table(rows=1, cols=1)
    sug_table.style = 'Table Grid'
    sug_row = sug_table.rows[0]
    sug_row.height = Cm(2)
    sug_row.cells[0].text = ''

    doc.add_paragraph()

    # Leave section
    p = doc.add_paragraph()
    r = p.add_run('Leave Application / Leave Taken')
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    for line in [
        '1.From (dd/mm/yyyy): _____________ to (dd/mm/yyyy) _____________ (  day(s))',
        '2. Reasons for taking leave: ________________________________________________________________',
        '3. Total number of days taken: ________________________________________________________________',
    ]:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run('I hereby declare that the information given above is correct.')
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    doc.add_paragraph()

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.style = 'Table Grid'
    sig_table.autofit = False
    lc = sig_table.rows[0].cells[0]
    rc = sig_table.rows[0].cells[1]
    lc.width = Inches(3)
    rc.width = Inches(3)
    lp = lc.paragraphs[0]
    lr = lp.add_run('Signature: _______________________')
    lr.font.size = Pt(11)
    lr.font.name = 'Times New Roman'
    rp = rc.paragraphs[0]
    rr = rp.add_run(f'Date (dd/mm/yyyy):   {trainee_date}')
    rr.font.size = Pt(11)
    rr.font.name = 'Times New Roman'
    # Remove borders
    for cell in [lc, rc]:
        set_cell_border(cell, top='none', left='none', bottom='none', right='none')

    doc.add_paragraph()
    doc.add_paragraph()

    # Supervisor section
    p = doc.add_paragraph()
    r = p.add_run('Endorsement by the Company Supervisor:')
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run('The above is a true record of activities taken by the trainee in the captioned week.')
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

    doc.add_paragraph()

    sup_table = doc.add_table(rows=5, cols=2)
    sup_table.style = 'Table Grid'
    sup_data = [
        ('Signature of Supervisor:', ''),
        ('Name of Supervisor:', 'Ardyawati Binti Zainon'),
        ('Date (dd/mm/yyyy):', supervisor_date),
        ('Email:', 'ardyawati@engsheng.com'),
        ('Mobile / Office Contact No.:', '016-5227316'),
    ]
    for i, (label, val) in enumerate(sup_data):
        lc = sup_table.rows[i].cells[0]
        rc = sup_table.rows[i].cells[1]
        lc.width = Inches(2.5)
        rc.width = Inches(3.5)
        if i == 0:
            lc.paragraphs[0].paragraph_format.space_before = Pt(20)
            lc.paragraphs[0].paragraph_format.space_after = Pt(20)
            rc.paragraphs[0].paragraph_format.space_before = Pt(20)
            rc.paragraphs[0].paragraph_format.space_after = Pt(20)
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.size = Pt(11)
        lr.font.name = 'Times New Roman'
        rp = rc.paragraphs[0]
        rr = rp.add_run(val)
        rr.font.size = Pt(11)
        rr.font.name = 'Times New Roman'

    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run('Company Stamp:')
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'


# ---- JUNE DATA ----
june_weeks = [
    (1, [
        "Monitored the WhatsApp lorry assignment bot following its live deployment at the end of May. Observed the system handling real daily delivery order (DO) uploads from the logistics team, verified that assignment outputs and trip manifests were being delivered correctly to planners, and collected feedback from lorry coordinators on usability and accuracy of assignments.",
        "Investigated early post-deployment issues flagged by the operations team, including cases where delivery orders with future-dated remarks were being incorrectly filtered out before the user-selected trip day, and cases where the daily assignment log was not persisting correctly between bot sessions. Identified root causes in the remarks day parser and session state handling.",
        "Reviewed the complete assignment rule set against real operational data from the first week of live usage, cross-referencing assigned lorries against planner expectations. Documented a shortlist of rule inconsistencies and edge cases to be addressed in subsequent development iterations, covering route boundary enforcement, lorry capacity utilisation thresholds, and preferred-lorry override behaviour.",
    ]),
    (2, [
        "Added a route-direction guard to all DO consolidation and overflow assignment passes, preventing delivery orders from geographically opposite directions from being grouped onto the same lorry. This resolved a recurring complaint from lorry coordinators where a single lorry was being assigned stops spanning both the north and south corridors of the Klang Valley, causing excessive backtracking and longer trip durations.",
        "Fixed the lorry utilisation calculation to correctly enforce an 80% load threshold before triggering a LORRY NAIK (lorry upgrade) split. Corrected the fill-to-80% pass to ensure that partial loads are consolidated up to the utilisation target before a second lorry is assigned for the overflow, preventing under-loaded lorries from being dispatched unnecessarily.",
        "Implemented earliest-date-first prioritisation across all assignment passes, ensuring that delivery orders with earlier scheduled dates are always processed and assigned before later ones when the lorry fleet is at capacity. This prevents newer DOs from displacing older outstanding ones during high-volume days.",
        "Strengthened the urban route guard to enforce corridor-group-only cross-route mixing for Klang Valley deliveries. Routes belonging to different geographic corridor groups are now strictly blocked from being merged onto the same lorry, even when both routes are classified as urban, closing a gap where KL and Selangor routes from opposite ends of the city were being co-assigned.",
        "Fixed the Tomorrow button flow to correctly include Saturday as a selectable working day while skipping Sunday only, in line with the company's Monday-to-Saturday operating schedule.",
    ]),
    (3, [
        "Centralised all assignment business-rule constants into a dedicated assignment_config.py module and created a companion ASSIGNMENT_RULES.md file that is loaded at every system startup as the authoritative rules reference. This separates operational parameters from application logic, allowing rules to be updated without modifying core engine code, and gives the operations team a single readable document to review and approve rule changes.",
        "Introduced data-driven preferred lorry configuration, replacing hardcoded lorry-to-route mappings with a table-driven lookup. Added support for REMARKS FIELD tonnage caps, allowing delivery orders to specify a maximum lorry size via free-text remarks, with the parser handling variable phrasing such as 'BELOW OR 5 TON LORRY' and similar natural-language variants.",
        "Refactored the assignment merge logic to prioritise route matching first, then same-city consolidation, and finally same-state grouping by nearest longitude. Added support for same-route-prefix sub-buckets to be pre-merged before cross-route grouping, resolving a split issue where the KV01A route was being divided across two lorries instead of consolidated.",
        "Integrated OSRM road-distance routing using the OpenStreetMap road network as the primary stop-ordering method, with haversine straight-line distance as a fallback when OSRM is unavailable. Replaced the previous latitude-sweep stop ordering with a greedy nearest-neighbour algorithm that sequences manifest stops from the depot outward by actual road travel distance, improving route efficiency and reducing total trip distance.",
        "Enhanced the trip manifest to include a RETURN TO HQ row at the end of each lorry's stop list and a round-trip estimated time of arrival (ETA) summary, giving drivers a clearer picture of the expected total journey time and confirming the final return leg of each trip.",
        "Fixed multiple geographic stop-ordering bugs in the trip manifest, replacing incorrect longitude-only sorting with a principal-axis geographic sweep that correctly sequences stops from HQ outward in the actual direction of travel, avoiding the zigzag patterns produced by the previous one-dimensional sort.",
        "Resolved a state-compatibility exclusion bug in the NO_ELIGIBLE_LORRY path where multi-state route buckets were incorrectly filtering out compatible lorries. Fixed lorry preferred-owner logic to treat owner assignment as a hard constraint rather than a soft hint, falling back to other lorries only when all owner-lorries are at capacity.",
        "Wrote and published a comprehensive Operator and Maintenance Manual covering the full system architecture, WhatsApp user menu flow, all assignment rules, master data file formats, external service dependencies, environment variable configuration, restart procedures, maintenance instructions, troubleshooting steps, and quick reference cards for both operators and the maintenance engineer.",
    ]),
    (4, [
        "Fixed preferred lorry configuration for specific routes: BPE9788 is now reserved strictly for the PH (Pahang) route, the ABI lorry is configured as the fallback for KV14A, and the TR02 route strict assignment is enforced correctly. These corrections were identified through planner feedback during live operations where the wrong lorries were being suggested for familiar routes.",
        "Fixed two lorry-contamination bugs that were causing outstation-assigned lorries to appear in urban route assignments. The first bug allowed a lorry with an existing outstation booking to be selected for a Klang Valley DO when its outstation route overlapped with an urban route prefix. The second bug occurred during the overflow reassignment pass, where lorries already committed to outstation trips were incorrectly considered available for urban overflow.",
        "Implemented automatic splitting of oversized single delivery orders that exceed the maximum capacity of any single lorry. Previously, a DO that was too heavy for even the largest available lorry would be flagged as NO_LORRY with no resolution. The system now splits such DOs across multiple lorries proportionally, ensuring all items are assigned and nothing is left undelivered due to size constraints.",
    ]),
]

# ---- JULY DATA ----
july_weeks = [
    (1, [
        "Conducted structured user acceptance testing (UAT) sessions with the logistics planning team and lorry coordinators, running the deployed system against a full week of historical daily DO data to validate assignment accuracy and trip manifest correctness. Documented all cases where the system output differed from what an experienced planner would manually assign, and categorised them by root cause for targeted rule improvements.",
        "Analysed patterns in the NO_LORRY and NO_ELIGIBLE_LORRY assignment failures collected during UAT and live operations, identifying routes and weight combinations that consistently fall outside the current fleet capacity model. Proposed and began implementing an extended fallback assignment strategy that considers partial-day lorry availability and shared-load arrangements between routes as additional resolution options.",
        "Refined the remarks parsing engine to handle a broader range of free-text delivery restriction phrases encountered in real operational data, including bilingual and mixed-language entries combining Malay and English within the same remarks field. Extended the parser's negation detection to correctly handle compound clauses such as 'JANGAN HANTAR ISNIN DAN SELASA' without misclassifying the negation scope.",
    ]),
    (2, [
        "Developed a daily operations summary report feature that generates a concise WhatsApp-formatted text summary at the end of each assignment session, listing total DOs assigned, lorries utilised, average load percentage per lorry, and any unresolved NO_LORRY items. This gives the logistics manager a one-tap end-of-day snapshot without needing to open the exported Excel files.",
        "Prepared internal handover documentation covering the system's codebase structure, deployment configuration on the Windows server, ngrok tunnel management, Meta Cloud API token renewal procedure, and common troubleshooting scenarios. The documentation is written to enable a non-developer operations staff member to perform routine maintenance and restart the system independently after the internship period ends.",
    ]),
]

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# ---- JUNE REPORT ----
build_report(doc, 'JUNE/2026', june_weeks, '30/6/2026', '30/06/2026')

doc.add_page_break()

# ---- JULY REPORT ----
build_report(doc, 'JULY/2026', july_weeks, '14/7/2026', '14/07/2026')

doc.save('/home/user/DO_bot/June_July_2026_REPORT_OOI_YU_ZHEN.docx')
print("Done.")
